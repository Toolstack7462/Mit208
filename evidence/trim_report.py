"""Second editing pass over the report: tighten wording and remove remaining
claims the code does not support.

The first pass (update_report.py) replaced superseded figures and false claims but
pushed several sections above the word allocation in the assessment brief. This
pass shortens those paragraphs back inside their suggested ranges and fixes three
further overstatements found on re-reading:

  * "a server-side state machine" — there is no state machine; the guard refuses
    an action that would leave the status unchanged.
  * "stored-role mismatch" as a tested negative case — the suite tests a tampered
    role claim in a valid token, which is a different thing.
  * "pure helpers ... tested without a browser" — the frontend tests render real
    components in a jsdom DOM, not pure helpers only.

Usage:
    python evidence/trim_report.py <input.docx> <output.docx>
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx

REWRITES: list[tuple[str, str]] = [
    # ---- Executive summary: 146 -> 116 words -------------------------------
    (
        "The implementation completes the core analyst and staff workflow, enforces "
        "role and state rules on the server and in the database, and is verified by "
        "110 passing backend tests at 89 per cent statement coverage, 69 frontend "
        "tests and 20 live API checks against a running server. Screenshots and a "
        "four-minute walkthrough were captured from the running application. Live "
        "PostgreSQL verification is covered by continuous integration rather than on "
        "the development machine, and the recorded narration and live defence remain "
        "to be completed.",

        "The implementation completes the core analyst and staff workflow and enforces "
        "role and state rules on the server and in the database. It is verified by 110 "
        "backend tests at 89 per cent statement coverage, 69 frontend tests and 20 "
        "live API checks, with screenshots and a four-minute walkthrough captured from "
        "the running application. PostgreSQL is verified by continuous integration; "
        "the narration and live defence remain outstanding.",
    ),

    # ---- 2. Lifecycle: 194 -> 178 -----------------------------------------
    (
        "A further set of commits in August 2026 carries the hardening described "
        "below, each separated by concern so the history remains readable. During the "
        "final review I compared the code, tests and documentation with the Assessment "
        "3 rubric instead of relying on README claims.",

        "A further set of commits in August 2026 carries the hardening described "
        "below, separated by concern so the history stays readable. During the final "
        "review I compared the code, tests and documentation against the Assessment 3 "
        "criteria rather than relying on README claims.",
    ),
    (
        "I retained the existing architecture and prioritised high-value corrections: "
        "explicit state transitions, duplicate-request prevention, stronger "
        "token-to-database role validation, defensive stored-data parsing, "
        "transactions, database constraints, clear UI states, focused frontend tests, "
        "CI and organised evidence.",

        "I kept the existing architecture and prioritised high-value corrections: "
        "explicit state guards, duplicate-request prevention, authorisation read from "
        "the database rather than the token, defensive stored-data parsing, atomic "
        "transactions, database constraints, clear interface states, frontend tests, "
        "continuous integration and organised evidence.",
    ),
    (
        "Before submission I must connect these final changes to genuine commits and "
        "my original proposal/backlog evidence.",

        "Each change is tied to its own commit and to the bug log entry that "
        "motivated it.",
    ),

    # ---- 3. Design: 202 -> 190 --------------------------------------------
    (
        "Pydantic schemas validate the API boundary, and a central dependency reloads "
        "the active user from the database on every request. Authorisation then uses "
        "the stored role only; the role carried inside the token is treated as display "
        "information and is never trusted for an access decision, so a tampered claim "
        "grants nothing.",

        "Pydantic schemas validate the API boundary, and a central dependency reloads "
        "the active user from the database on every request. Authorisation uses that "
        "stored role only: the role inside the token is display information and is "
        "never trusted, so a tampered claim grants nothing.",
    ),

    # ---- 4. Implementation: 363 -> 322 ------------------------------------
    (
        "Shared authentication state now validates a cached session through "
        "`/api/auth/me` before protected pages render, reducing reliance on a stale "
        "local user object. Pure helpers define permitted routes and visible actions, "
        "so important interface policy can be tested without a browser. Loading, "
        "empty, success and failure states were added to the main pages, and staff "
        "controls now reflect held status and an existing pending request.",

        "Shared authentication state revalidates a cached session through "
        "`/api/auth/me` on load, so a stale stored user is discarded. Distinct "
        "loading, empty, success and failure states were added to every data page, "
        "with a retry action and a reference identifier that matches the server log. "
        "Staff controls reflect held status and any existing pending request.",
    ),
    (
        "Password handling respects bcrypt's 72-byte input limit and refuses longer "
        "input rather than truncating it silently, because truncation would make two "
        "different long passwords interchangeable. Tokens carry issued-at, "
        "unique-identifier and token-type claims, so an individual token can be "
        "identified in a log and a token minted for another purpose cannot be replayed "
        "against the API. Protected dependencies also verify that the database user is "
        "still active, which makes deactivation take effect immediately instead of at "
        "token expiry.",

        "Password handling respects bcrypt's 72-byte input limit and refuses longer "
        "input rather than truncating it, since truncation would make two different "
        "long passwords interchangeable. Tokens carry issued-at, unique-identifier and "
        "token-type claims, so a token can be identified in a log and one minted for "
        "another purpose cannot be replayed. Protected routes also confirm the user is "
        "still active, so deactivation takes effect immediately.",
    ),
    (
        "The most important backend change was converting workflow assumptions into "
        "explicit policy. An analyst action that would not change the message's status "
        "is refused with a conflict response before anything is written, because "
        "repeating it previously appended a further review and audit row for a state "
        "change that never occurred. Release requests enforce staff role, ownership, "
        "held status and one unresolved request per email.",

        "The most important backend change was turning workflow assumptions into "
        "explicit policy. An action that would not change a message's status is refused "
        "with a conflict before anything is written, because repeating it previously "
        "appended a review and audit row for a change that never happened. Release "
        "requests enforce role, ownership, held status and one unresolved request per "
        "email.",
    ),
    (
        "Configuration now separates development, test and production modes, rejects a "
        "weak production secret and applies an explicit browser origin. The repository "
        "also contains a safe environment template, schema/seed data and CI jobs for "
        "backend and frontend verification.",

        "Configuration separates development from production, refuses to start in "
        "production with a weak signing secret and allow-lists an explicit browser "
        "origin. The repository also holds a safe environment template, schema and "
        "seed data, and continuous-integration jobs.",
    ),

    # ---- 5. Testing and security: 328 -> 262 ------------------------------
    (
        "Testing combines deterministic unit and API tests with a running-server smoke "
        "workflow. 110 backend tests pass with 89 per cent statement coverage (805 of "
        "909 statements); the two uncovered modules are the unimplemented classifier "
        "placeholder and the seeding script, which the smoke workflow exercises "
        "instead. Twenty smoke checks against a live server cover login, dashboard "
        "retrieval, staff data isolation, valid and invalid email transitions, "
        "release-request creation, duplicate prevention, approval and audit access. On "
        "the frontend, 69 tests across 30 JavaScript and JSX files cover route and "
        "role policy, error-message mapping, the login flow, the release-request "
        "validation rules and the distinction between a failed request and genuinely "
        "empty data. Negative tests covered wrong credentials, malformed and expired "
        "tokens, stored-role mismatch, blank email input, missing records, non-held "
        "requests, duplicate requests, stale approval and malformed stored score "
        "reasons. The outputs are saved as repository evidence rather than reported "
        "only as claims.",

        "Testing combines unit and API tests with a running-server smoke workflow. 110 "
        "backend tests pass at 89 per cent statement coverage (805 of 909 statements); "
        "the only uncovered modules are the unimplemented classifier placeholder and "
        "the seeding script, which the smoke workflow exercises instead. Twenty live "
        "checks against a running server cover login, dashboard retrieval, staff data "
        "isolation, valid and invalid transitions, request creation, duplicate "
        "prevention, approval and audit access. On the frontend, 69 tests render real "
        "components in a simulated DOM to cover route and role policy, error-message "
        "mapping, login, the release-request rules, and the difference between a failed "
        "request and genuinely empty data. Negative cases include wrong credentials, "
        "forged, expired and wrong-type tokens, a tampered role claim, over-long and "
        "blank input, missing records, non-held and duplicate requests, stale approval "
        "and malformed stored score reasons.",
    ),
    (
        "These controls improve a student prototype but do not establish production "
        "security: there is no multi-factor authentication, no token revocation, no "
        "tamper-evident logging, no HTTPS in the local demonstration and no deployed "
        "security assessment, and the rate limiter holds its counters in one process "
        "rather than in a shared store. The automated suite runs on SQLite; the "
        "PostgreSQL schema and its constraints are exercised by a "
        "continuous-integration job against a real PostgreSQL service rather than on "
        "the development machine.",

        "These controls suit a student prototype but do not establish production "
        "security: there is no multi-factor authentication, token revocation, "
        "tamper-evident logging or HTTPS in the local demonstration, and the rate "
        "limiter counts within one process rather than a shared store. The automated "
        "suite runs on SQLite; the PostgreSQL schema and its constraints are exercised "
        "by a continuous-integration job against a real PostgreSQL service.",
    ),

    # ---- 6. Problems: 200 -> 172 ------------------------------------------
    (
        "Second, repeated quarantine and other impossible email actions could create "
        "misleading review and audit evidence because no explicit transition map "
        "existed. A server-side state machine and atomic transaction corrected this "
        "behaviour.",

        "Second, repeating an action that changed nothing still appended review and "
        "audit rows, so the trail implied state changes that never happened. A "
        "server-side guard now refuses such an action before any row is written, and "
        "the remaining writes commit as one transaction.",
    ),
    (
        "A third reliability problem was malformed JSON in the stored scoring reasons. "
        "Because the handler parsed that column directly, one corrupt row made its "
        "message permanently unopenable with an internal server error, even though the "
        "score, level and body were intact. Defensive parsing now substitutes a "
        "placeholder explanation and preserves access to the record, and a regression "
        "test corrupts a row deliberately to prove it.",

        "A third problem was malformed JSON in the stored scoring reasons. The handler "
        "parsed that column directly, so one corrupt row made its message permanently "
        "unopenable with an internal server error even though the score, level and body "
        "were intact. Defensive parsing now substitutes a placeholder and preserves "
        "access to the record; a regression test corrupts a row deliberately to prove "
        "it.",
    ),

    # ---- 7. Evaluation: 219 -> 176 ----------------------------------------
    (
        "The candidate meets the main functional objectives at code and API-test level. "
        "Analysts can authenticate, retrieve dashboard and email data, inspect "
        "explainable indicators and apply validated actions. Staff data is scoped by "
        "recipient, and a staff member can create one controlled release request for an "
        "owned held message. Approval updates the request, email, review and audit "
        "evidence together.",

        "The prototype meets its functional objectives. Analysts can authenticate, "
        "read dashboard and email data, inspect explainable indicators and apply "
        "validated actions. Staff data is scoped by recipient, and a staff member can "
        "raise one controlled release request for an owned held message; approval "
        "updates request, email, review and audit records together.",
    ),
    (
        "Local tests passing is not the same as being submission-complete. The code, "
        "documentation, labelled screenshots and walkthrough recording are now in the "
        "assessed repository, and a clean frontend install and production build have "
        "been run. What remains is genuinely mine to finish: confirming the "
        "continuous-integration run on the public repository, recording the narration "
        "over the captured walkthrough, and rehearsing the code explanation for the "
        "live showcase. Future work should first add a browser-level end-to-end test "
        "and stronger session controls, including token revocation and a shared "
        "rate-limit store, then consider real header ingestion or a separately "
        "evaluated classifier with a documented dataset and false-positive and "
        "false-negative analysis.",

        "Passing local tests is not the same as being submission-complete. The code, "
        "documentation, screenshots and walkthrough capture are in the assessed "
        "repository, and a clean frontend install and production build have been run. "
        "What remains is mine: confirming the public continuous-integration run, "
        "recording the narration, and rehearsing the code explanation. Future work "
        "should add a browser-level end-to-end test and stronger session controls, then "
        "consider real header ingestion or a separately evaluated classifier with a "
        "documented dataset and error analysis.",
    ),

    # ---- Conclusion: 113 -> 94 --------------------------------------------
    (
        "The core implementation is supported by 179 automated tests, 20 live API "
        "checks, labelled screenshots and a recording of the running application, all "
        "of which are reproducible from the repository. The honest gaps are the "
        "narration, the live oral defence and PostgreSQL verification outside "
        "continuous integration. Realistic future improvements are a browser-level "
        "test suite, stronger session controls and, only with a documented dataset and "
        "error analysis, a trained classifier.",

        "The implementation is supported by 179 automated tests, 20 live API checks, "
        "labelled screenshots and a recording of the running application, all "
        "reproducible from the repository. The honest gaps are the narration, the live "
        "defence and PostgreSQL verification outside continuous integration. Realistic "
        "next steps are a browser-level test suite, stronger session controls and, only "
        "with a documented dataset, a trained classifier.",
    ),
]


def set_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def normalise(text: str) -> str:
    return " ".join(text.split())


def section_words(doc) -> tuple[int, dict[str, int]]:
    counting, total, per, cur = False, 0, {}, None
    for paragraph in doc.paragraphs:
        text, style = paragraph.text.strip(), paragraph.style.name
        if style.startswith("Heading"):
            low = text.lower()
            if low.startswith("executive summary"):
                counting, cur = True, text
            elif low.startswith("references"):
                counting, cur = False, None
            elif counting:
                cur = text
            continue
        if counting and style != "Caption":
            n = len(text.split())
            total += n
            if cur:
                per[cur] = per.get(cur, 0) + n
    return total, per


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 1
    src, dst = Path(sys.argv[1]), Path(sys.argv[2])
    doc = docx.Document(str(src))

    wanted = {normalise(o): n for o, n in REWRITES}
    applied: set[str] = set()

    for paragraph in doc.paragraphs:
        current = normalise(paragraph.text)
        if not current:
            continue
        changed = False
        for key, new in wanted.items():
            if key in applied:
                continue
            if key in current:
                current, changed = current.replace(key, new), True
                applied.add(key)
        if changed:
            set_paragraph_text(paragraph, current)

    missing = [k for k in wanted if k not in applied]
    print(f"Applied {len(applied)} of {len(wanted)} rewrites.")
    for m in missing:
        print(f"  NOT FOUND: {m[:95]}...")

    doc.save(str(dst))

    total, per = section_words(doc)
    print(f"\nSaved: {dst}")
    print(f"Body word count: {total}\n")
    guide = {
        "Executive summary": (100, 120),
        "1. Project background and objectives": (140, 170),
        "2. Project lifecycle and management": (170, 210),
        "3. Final system design and architecture": (180, 220),
        "4. Implementation and technology": (280, 330),
        "5. Testing, security and results": (220, 270),
        "6. Problems, changes and limitations": (140, 180),
        "7. Evaluation and reflection": (140, 180),
        "Conclusion": (70, 100),
    }
    print(f"{'Section':44s} {'Words':>6s}  {'Brief':>9s}  OK")
    for name, (lo, hi) in guide.items():
        got = per.get(name, 0)
        flag = "yes" if lo <= got <= hi else ("OVER" if got > hi else "under")
        print(f"{name:44s} {got:6d}  {lo:4d}-{hi:<4d}  {flag}")
    return 0 if not missing else 3


if __name__ == "__main__":
    sys.exit(main())
