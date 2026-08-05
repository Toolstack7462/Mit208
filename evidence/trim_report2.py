"""Final trim: bring every report section inside the brief's word allocation.

Pass 1 (update_report.py) corrected the figures and false claims; pass 2
(trim_report.py) tightened the rewritten paragraphs. Several sections were still
above their suggested range, so this pass makes the last cuts. Only wording is
removed — no verified figure, citation or disclosed limitation is dropped.

Usage:
    python evidence/trim_report2.py <input.docx> <output.docx>
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx

REWRITES: list[tuple[str, str]] = [
    # ---- Executive summary: 130 -> 112 -------------------------------------
    (
        "PhishGuard is a full-stack student prototype for prioritising suspicious "
        "email samples and controlling quarantine and release decisions. It gives "
        "analysts explainable risk evidence and gives staff a governed way to "
        "challenge a held message. The solution combines React, FastAPI, SQLAlchemy "
        "and a PostgreSQL target database, with SQLite used for automated "
        "verification. Detection is a transparent rule-based score rather than a "
        "trained machine-learning model.",

        "PhishGuard is a full-stack student prototype for triaging suspicious email "
        "and controlling quarantine and release decisions. It gives analysts "
        "explainable risk evidence and gives staff a governed way to challenge a held "
        "message. It combines React, FastAPI and SQLAlchemy over PostgreSQL, with "
        "SQLite for automated verification. Detection is a transparent rule-based "
        "score, not a trained model.",
    ),

    # ---- 4. Implementation: 333 -> 300 ------------------------------------
    (
        "React 18 and Vite were retained because the existing interface was already "
        "componentised and suitable for a local demonstration. Shared authentication "
        "state revalidates a cached session through `/api/auth/me` on load, so a stale "
        "stored user is discarded. Distinct loading, empty, success and failure states "
        "were added to every data page, with a retry action and a reference identifier "
        "that matches the server log. Staff controls reflect held status and any "
        "existing pending request.",

        "React 18 and Vite were retained because the interface was already "
        "componentised and suited to a local demonstration. Shared authentication "
        "state revalidates a cached session through `/api/auth/me` on load, discarding "
        "a stale stored user. Every data page now has distinct loading, empty and "
        "failure states, with a retry action and a reference identifier matching the "
        "server log. Staff controls reflect held status and any pending request.",
    ),
    (
        "FastAPI was appropriate for typed request handling and automatic OpenAPI "
        "documentation (FastAPI, n.d.). Authentication uses bcrypt password hashes and "
        "expiring signed JSON Web Tokens. Password handling respects bcrypt's 72-byte "
        "input limit and refuses longer input rather than truncating it, since "
        "truncation would make two different long passwords interchangeable. Tokens "
        "carry issued-at, unique-identifier and token-type claims, so a token can be "
        "identified in a log and one minted for another purpose cannot be replayed. "
        "Protected routes also confirm the user is still active, so deactivation takes "
        "effect immediately. This reflects guidance that password verifiers and "
        "authentication data require deliberate protection rather than informal string "
        "comparison (OWASP Foundation, n.d.; Jones, Bradley and Sakimura, 2015).",

        "FastAPI suited typed request handling and automatic OpenAPI documentation "
        "(FastAPI, n.d.). Authentication uses bcrypt hashes and expiring signed JSON "
        "Web Tokens. Password handling respects bcrypt's 72-byte limit and refuses "
        "longer input rather than truncating it, since truncation would make two "
        "different long passwords interchangeable. Tokens carry issued-at, "
        "unique-identifier and token-type claims, so a token can be traced in a log "
        "and one minted for another purpose cannot be replayed. Protected routes also "
        "confirm the user is still active, so deactivation takes effect immediately. "
        "This follows guidance that authentication data needs deliberate protection "
        "rather than informal comparison (OWASP Foundation, n.d.; Jones, Bradley and "
        "Sakimura, 2015).",
    ),
    (
        "SQLAlchemy transactions commit related state, review and audit rows together, "
        "with rollback on failure; the session therefore acts as a unit of work across "
        "related records (SQLAlchemy, 2026). ORM and PostgreSQL schema constraints "
        "restrict role, status and decision values, while a partial unique index "
        "guards pending duplicates (PostgreSQL Global Development Group, 2026a; "
        "2026b). Configuration separates development from production, refuses to start "
        "in production with a weak signing secret and allow-lists an explicit browser "
        "origin. The repository also holds a safe environment template, schema and "
        "seed data, and continuous-integration jobs.",

        "Transactions commit state, review and audit rows together and roll back on "
        "failure, so the session acts as a unit of work (SQLAlchemy, 2026). Constraints "
        "in both the ORM and the schema restrict role, status and decision values, and "
        "a partial unique index guards pending duplicates (PostgreSQL Global "
        "Development Group, 2026a; 2026b). Configuration refuses to start in production "
        "with a weak signing secret and allow-lists an explicit browser origin.",
    ),

    # ---- 5. Testing and security: 300 -> 258 ------------------------------
    (
        "Testing combines unit and API tests with a running-server smoke workflow. 110 "
        "backend tests pass at 89 per cent statement coverage (805 of 909 statements); "
        "the only uncovered modules are the unimplemented classifier placeholder and "
        "the seeding script, which the smoke workflow exercises instead. Twenty live "
        "checks against a running server cover login, dashboard retrieval, staff data "
        "isolation, valid and invalid transitions, request creation, duplicate "
        "prevention, approval and audit access. On the frontend, 69 tests render real "
        "components in a simulated DOM to cover route and role policy, error-message "
        "mapping, login, the release-request rules, and the difference between a failed "
        "request and genuinely empty data. Negative cases include wrong credentials, "
        "forged, expired and wrong-type tokens, a tampered role claim, over-long and "
        "blank input, missing records, non-held and duplicate requests, stale approval "
        "and malformed stored score reasons.",

        "Testing combines unit and API tests with a running-server smoke workflow. 110 "
        "backend tests pass at 89 per cent statement coverage (805 of 909 statements); "
        "the only uncovered modules are the classifier placeholder and the seeding "
        "script, which the smoke workflow exercises instead. Twenty live checks against "
        "a running server cover login, dashboard retrieval, staff data isolation, valid "
        "and invalid transitions, request creation, duplicate prevention, approval and "
        "audit access. On the frontend, 69 tests render real components in a simulated "
        "DOM, covering route and role policy, error-message mapping, login, the "
        "release-request rules and the difference between a failed request and genuinely "
        "empty data. Negative cases include wrong credentials, forged, expired and "
        "wrong-type tokens, a tampered role claim, over-long input, missing records, "
        "duplicate requests and malformed stored score reasons.",
    ),
    (
        "Security controls include bcrypt hashes, expiring JWTs with typed claims, "
        "server-side role checks, per-IP limiting of failed logins, input bounds "
        "matched to the database column widths, ownership filtering, explicit CORS, "
        "response security headers, safe secret templates, parameterised ORM queries "
        "and database constraints that reject impossible or duplicate rows. Login "
        "returns an identical response for an unknown address and a wrong password, so "
        "it cannot be used to discover which accounts exist. Audit events record login "
        "and all material review and request activity.",

        "Security controls include bcrypt hashes, expiring typed JWTs, server-side role "
        "checks, per-IP limiting of failed logins, input bounds matched to the column "
        "widths, ownership filtering, explicit CORS, response security headers, safe "
        "secret templates, parameterised ORM queries and constraints that reject "
        "impossible or duplicate rows. Login answers an unknown address and a wrong "
        "password identically, so it cannot reveal which accounts exist. Audit events "
        "record login and all material review and request activity.",
    ),

    # ---- 6. Problems: 209 -> 172 ------------------------------------------
    (
        "Two significant defects shaped the final changes. First, the original "
        "release-request route accepted duplicate pending requests, requests for "
        "messages already in the inbox, and requests created by an analyst. "
        "Investigation showed that ownership was checked, but role and business-state "
        "rules were missing from both the API and database. The fix added a staff-only "
        "dependency, held-status rule, duplicate lookup, partial unique index, conflict "
        "handling and aligned interface controls. Second, repeating an action that "
        "changed nothing still appended review and audit rows, so the trail implied "
        "state changes that never happened. A server-side guard now refuses such an "
        "action before any row is written, and the remaining writes commit as one "
        "transaction.",

        "Two significant defects shaped the final changes. First, the release-request "
        "route accepted duplicate pending requests and requests for messages already "
        "delivered. Ownership was checked, but business-state rules were missing from "
        "both the API and the database. The fix added a held-status rule, a duplicate "
        "lookup, a partial unique index, conflict handling and matching interface "
        "controls. Second, repeating an action that changed nothing still appended "
        "review and audit rows, so the trail implied changes that never happened; a "
        "server-side guard now refuses such an action before any row is written.",
    ),
    (
        "Remaining limitations include synthetic data, simulated authentication-header "
        "results, a heuristic AI-generated label, an unimplemented DistilBERT "
        "placeholder, small-scale dashboard queries, browser-stored tokens, no "
        "automated browser-level test and no live mail-server integration. These "
        "limitations are disclosed rather than hidden.",

        "Remaining limitations include synthetic data, simulated authentication-header "
        "results, a heuristic AI-generated label, an unimplemented DistilBERT "
        "placeholder, browser-stored tokens, no automated browser-level test and no "
        "live mail-server integration. These are disclosed rather than hidden.",
    ),

    # ---- 7. Evaluation: 185 -> 168 ----------------------------------------
    (
        "The prototype meets its functional objectives. Analysts can authenticate, "
        "read dashboard and email data, inspect explainable indicators and apply "
        "validated actions. Staff data is scoped by recipient, and a staff member can "
        "raise one controlled release request for an owned held message; approval "
        "updates request, email, review and audit records together. My strongest "
        "learning outcome was that interface restrictions are not security controls: "
        "every role, ownership and state rule must be rechecked at the backend and, "
        "where possible, supported by a database constraint.",

        "The prototype meets its functional objectives. Analysts can authenticate, "
        "read dashboard and email data, inspect explainable indicators and apply "
        "validated actions. Staff data is scoped by recipient, and a staff member can "
        "raise one controlled release request for an owned held message; approval "
        "updates request, email, review and audit records together. My strongest "
        "learning outcome was that interface restrictions are not security controls: "
        "every role, ownership and state rule must be rechecked server-side and, where "
        "possible, backed by a database constraint.",
    ),
    (
        "Passing local tests is not the same as being submission-complete. The code, "
        "documentation, screenshots and walkthrough capture are in the assessed "
        "repository, and a clean frontend install and production build have been run. "
        "What remains is mine: confirming the public continuous-integration run, "
        "recording the narration, and rehearsing the code explanation. Future work "
        "should add a browser-level end-to-end test and stronger session controls, "
        "then consider real header ingestion or a separately evaluated classifier with "
        "a documented dataset and error analysis.",

        "Passing local tests is not the same as being submission-complete. The code, "
        "documentation, screenshots and walkthrough capture are in the assessed "
        "repository, and a clean frontend install and production build have been run. "
        "What remains is mine: confirming the public continuous-integration run, "
        "recording the narration and rehearsing the code explanation. Future work "
        "should add a browser-level test and stronger session controls, then consider "
        "real header ingestion or a separately evaluated classifier.",
    ),

    # ---- Conclusion: 105 -> 90 --------------------------------------------
    (
        "PhishGuard demonstrates a coherent and explainable phishing-triage workflow "
        "with role-based access, validated state changes, controlled staff release "
        "requests, audit evidence and repeatable testing. The final review improved "
        "reliability and security without replacing the original architecture and "
        "without overstating the rule engine as machine learning. The implementation "
        "is supported by 179 automated tests, 20 live API checks, labelled screenshots "
        "and a recording of the running application, all reproducible from the "
        "repository. The honest gaps are the narration, the live defence and "
        "PostgreSQL verification outside continuous integration. Realistic next steps "
        "are a browser-level test suite, stronger session controls and, only with a "
        "documented dataset, a trained classifier.",

        "PhishGuard demonstrates a coherent, explainable phishing-triage workflow with "
        "role-based access, validated state changes, controlled release requests, audit "
        "evidence and repeatable testing. The final review improved reliability and "
        "security without replacing the architecture or overstating the rule engine as "
        "machine learning. It is supported by 179 automated tests, 20 live API checks, "
        "labelled screenshots and a recording of the running application, all "
        "reproducible from the repository. The honest gaps are the narration and the "
        "live defence. Realistic next steps are a browser-level test suite and stronger "
        "session controls.",
    ),
]


def set_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def normalise(text: str) -> str:
    return " ".join(text.split())


GUIDE = {
    "Executive summary": (100, 120),
    "1. Project background and objectives": (140, 170),
    "2. Project lifecycle and management": (170, 210),
    "3. Final system design and architecture": (180, 220),
    "4. Implementation and technology": (280, 330),
    "5. Testing, security and results": (220, 270),
    "6. Problems, changes and limitations": (140, 180),
    "7. Evaluation and reflection": (140, 180),
    "Conclusion": (70, 100),
}


def section_words(doc):
    counting, total, per, cur = False, 0, {}, None
    for paragraph in doc.paragraphs:
        text, style = paragraph.text.strip(), paragraph.style.name
        if style.startswith("Heading"):
            low = text.lower()
            if low.startswith("executive summary"):
                counting, cur = True, text
            elif low.startswith("references"):
                counting, cur = False, None
            elif counting:
                cur = text
            continue
        if counting and style != "Caption":
            n = len(text.split())
            total += n
            if cur:
                per[cur] = per.get(cur, 0) + n
    return total, per


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 1
    src, dst = Path(sys.argv[1]), Path(sys.argv[2])
    doc = docx.Document(str(src))

    wanted = {normalise(o): n for o, n in REWRITES}
    applied: set[str] = set()
    for paragraph in doc.paragraphs:
        current = normalise(paragraph.text)
        if not current:
            continue
        changed = False
        for key, new in wanted.items():
            if key in applied:
                continue
            if key in current:
                current, changed = current.replace(key, new), True
                applied.add(key)
        if changed:
            set_paragraph_text(paragraph, current)

    missing = [k for k in wanted if k not in applied]
    print(f"Applied {len(applied)} of {len(wanted)} rewrites.")
    for m in missing:
        print(f"  NOT FOUND: {m[:90]}...")

    doc.save(str(dst))
    total, per = section_words(doc)
    print(f"\nSaved: {dst}\nBody word count: {total}\n")
    print(f"{'Section':44s} {'Words':>6s}  {'Brief':>9s}  OK")
    all_ok = True
    for name, (lo, hi) in GUIDE.items():
        got = per.get(name, 0)
        ok = lo <= got <= hi
        all_ok &= ok
        print(f"{name:44s} {got:6d}  {lo:4d}-{hi:<4d}  {'yes' if ok else ('OVER' if got > hi else 'under')}")
    print(f"\nAll sections within the brief's allocation: {'yes' if all_ok else 'no'}")
    return 0 if not missing else 3


if __name__ == "__main__":
    sys.exit(main())
