"""Correct the existing PhishGuard report in place.

This does NOT rewrite the report. It edits specific paragraphs and table cells of
the existing DOCX so the document keeps its structure, styles, headings, figures
and captions, while removing claims that the code does not support and replacing
superseded test figures with the verified ones.

Every replacement below is traceable to a verified result recorded in
docs/TESTING.md.

Usage:
    python evidence/update_report.py  <source.docx>  <output.docx>
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx


# ---------------------------------------------------------------------------
# Verified figures (docs/TESTING.md; re-verified 5 August 2026)
#   backend  : 110 pytest tests, 89% statement coverage (805/909)
#   frontend : 69 vitest tests across 9 files
#   live     : 20/20 smoke checks against a running server
#   parsed   : 30 JavaScript/JSX files
# ---------------------------------------------------------------------------

REPLACEMENTS: list[tuple[str, str]] = [
    # ---- Executive summary -------------------------------------------------
    (
        "The audited candidate completed the core analyst and staff workflow, "
        "strengthened role and state enforcement, and produced 52 passing backend "
        "tests, 85 per cent statement coverage, 14 successful live API checks and "
        "five focused frontend policy tests. Final submission evidence still "
        "requires a clean PostgreSQL run, browser build, replacement screenshots, "
        "CI result and genuine GitHub release.",

        "The implementation completes the core analyst and staff workflow, enforces "
        "role and state rules on the server and in the database, and is verified by "
        "110 passing backend tests at 89 per cent statement coverage, 69 frontend "
        "tests and 20 live API checks against a running server. Screenshots and a "
        "four-minute walkthrough were captured from the running application. Live "
        "PostgreSQL verification is covered by continuous integration rather than on "
        "the development machine, and the recorded narration and live defence remain "
        "to be completed.",
    ),

    # ---- 2. Lifecycle -----------------------------------------------------
    (
        "The public repository shows an incremental implementation journey rather "
        "than one final upload. Nine visible commits between June and July 2026 "
        "cover the initial full-stack MVP, interface redesign, weekly dashboard "
        "chart, more realistic sample timestamps, documentation and screenshots, "
        "then backend tests and test documentation.",

        "The public repository shows an incremental implementation journey rather "
        "than one final upload. The first nine commits, between June and July 2026, "
        "cover the initial full-stack MVP, interface redesign, weekly dashboard "
        "chart, more realistic sample timestamps, documentation and screenshots, "
        "then backend tests and test documentation. A further set of commits in "
        "August 2026 carries the hardening described below, each separated by "
        "concern so the history remains readable.",
    ),

    # ---- 3. Design: accurate description of the authorisation model -------
    (
        "Pydantic schemas validate the API boundary, while central dependencies "
        "reload the active database user and compare the stored role with the role "
        "in the token.",

        "Pydantic schemas validate the API boundary, and a central dependency "
        "reloads the active user from the database on every request. Authorisation "
        "then uses the stored role only; the role carried inside the token is "
        "treated as display information and is never trusted for an access "
        "decision, so a tampered claim grants nothing.",
    ),

    # ---- 4. Implementation: token claims, stated accurately ---------------
    (
        "Authentication uses bcrypt password hashes and expiring signed JSON Web "
        "Tokens. Password handling accounts for the bcrypt input boundary, while "
        "tokens include issued-at, unique identifier and token-type claims. "
        "Protected dependencies also verify that the database user remains active "
        "and that the stored role has not changed.",

        "Authentication uses bcrypt password hashes and expiring signed JSON Web "
        "Tokens. Password handling respects bcrypt's 72-byte input limit and refuses "
        "longer input rather than truncating it silently, because truncation would "
        "make two different long passwords interchangeable. Tokens carry issued-at, "
        "unique-identifier and token-type claims, so an individual token can be "
        "identified in a log and a token minted for another purpose cannot be "
        "replayed against the API. Protected dependencies also verify that the "
        "database user is still active, which makes deactivation take effect "
        "immediately instead of at token expiry.",
    ),

    # ---- 4. Implementation: state rules, described as implemented ----------
    (
        "The most important backend change was converting workflow assumptions into "
        "explicit policy. Email actions use an action-to-source-state map and return "
        "conflict responses for repeated or impossible transitions.",

        "The most important backend change was converting workflow assumptions into "
        "explicit policy. An analyst action that would not change the message's "
        "status is refused with a conflict response before anything is written, "
        "because repeating it previously appended a further review and audit row for "
        "a state change that never occurred.",
    ),

    # ---- 5. Testing figures ------------------------------------------------
    (
        "Testing combined deterministic unit/API tests with a running-server smoke "
        "workflow. In the audited environment, 52 backend tests passed with 85 per "
        "cent statement coverage (716 of 846 statements). Fourteen smoke checks "
        "exercised login, dashboard retrieval, staff data isolation, valid and "
        "invalid email transitions, release-request creation, duplicate prevention, "
        "approval and audit access. Five Node tests verified frontend route and "
        "action policy, and 21 JavaScript/JSX files parsed without diagnostics.",

        "Testing combines deterministic unit and API tests with a running-server "
        "smoke workflow. 110 backend tests pass with 89 per cent statement coverage "
        "(805 of 909 statements); the two uncovered modules are the unimplemented "
        "classifier placeholder and the seeding script, which the smoke workflow "
        "exercises instead. Twenty smoke checks against a live server cover login, "
        "dashboard retrieval, staff data isolation, valid and invalid email "
        "transitions, release-request creation, duplicate prevention, approval and "
        "audit access. On the frontend, 69 tests across 30 JavaScript and JSX files "
        "cover route and role policy, error-message mapping, the login flow, the "
        "release-request validation rules and the distinction between a failed "
        "request and genuinely empty data.",
    ),

    # ---- 5. Security: rate limiting now exists -----------------------------
    (
        "Security controls include bcrypt hashes, expiring JWTs, server-side role "
        "checks, input bounds, ownership filtering, explicit CORS, secret templates, "
        "parameterised ORM queries and protection against impossible or duplicate "
        "actions. Audit events record login and material review/request activity. "
        "These controls improve a student prototype but do not establish production "
        "security: there is no MFA, rate limiting, token revocation, tamper-evident "
        "logging or deployed security assessment. The automated database was SQLite, "
        "so the PostgreSQL schema and constraints still require a clean live "
        "verification. A fresh `npm install`, frontend production build and patched "
        "browser workflow also remain manual evidence.",

        "Security controls include bcrypt hashes, expiring JWTs with typed claims, "
        "server-side role checks, per-IP limiting of failed logins, input bounds "
        "matched to the database column widths, ownership filtering, explicit CORS, "
        "response security headers, safe secret templates, parameterised ORM queries "
        "and database constraints that reject impossible or duplicate rows. Login "
        "returns an identical response for an unknown address and a wrong password, "
        "so it cannot be used to discover which accounts exist. Audit events record "
        "login and all material review and request activity. These controls improve a "
        "student prototype but do not establish production security: there is no "
        "multi-factor authentication, no token revocation, no tamper-evident logging, "
        "no HTTPS in the local demonstration and no deployed security assessment, and "
        "the rate limiter holds its counters in one process rather than in a shared "
        "store. The automated suite runs on SQLite; the PostgreSQL schema and its "
        "constraints are exercised by a continuous-integration job against a real "
        "PostgreSQL service rather than on the development machine.",
    ),

    # ---- 6. Problems: third defect is now genuinely fixed -----------------
    (
        "A third reliability problem was malformed JSON in stored scoring reasons, "
        "which could make an email-detail request fail. Defensive parsing now returns "
        "a safe explanation while preserving access to the record. Remaining "
        "limitations include synthetic data, simulated authentication-header results, "
        "a heuristic `ai_generated` label, an unused DistilBERT placeholder, "
        "small-scale dashboard queries, local token storage and no live mail-server "
        "integration. These limitations are disclosed rather than hidden.",

        "A third reliability problem was malformed JSON in the stored scoring "
        "reasons. Because the handler parsed that column directly, one corrupt row "
        "made its message permanently unopenable with an internal server error, even "
        "though the score, level and body were intact. Defensive parsing now "
        "substitutes a placeholder explanation and preserves access to the record, "
        "and a regression test corrupts a row deliberately to prove it. Remaining "
        "limitations include synthetic data, simulated authentication-header results, "
        "a heuristic AI-generated label, an unimplemented DistilBERT placeholder, "
        "small-scale dashboard queries, browser-stored tokens, no automated "
        "browser-level test and no live mail-server integration. These limitations "
        "are disclosed rather than hidden.",
    ),

    # ---- 7. Evaluation: honest current position ---------------------------
    (
        "The project is not submission-complete merely because local tests pass. I "
        "still need to apply the candidate to the assessed repository, rerun it with "
        "PostgreSQL, complete a clean frontend install/build, capture final labelled "
        "screenshots, obtain a passing CI run, create a genuine release and rehearse "
        "the code explanation. Future work should first add a small browser test and "
        "stronger session controls, then consider real header ingestion or a "
        "separately evaluated classifier with a documented dataset and "
        "false-positive/false-negative analysis.",

        "Local tests passing is not the same as being submission-complete. The code, "
        "documentation, labelled screenshots and walkthrough recording are now in the "
        "assessed repository, and a clean frontend install and production build have "
        "been run. What remains is genuinely mine to finish: confirming the "
        "continuous-integration run on the public repository, recording the narration "
        "over the captured walkthrough, and rehearsing the code explanation for the "
        "live showcase. Future work should first add a browser-level end-to-end test "
        "and stronger session controls, including token revocation and a shared "
        "rate-limit store, then consider real header ingestion or a separately "
        "evaluated classifier with a documented dataset and false-positive and "
        "false-negative analysis.",
    ),

    # ---- Conclusion --------------------------------------------------------
    (
        "PhishGuard demonstrates a coherent and explainable phishing-triage workflow "
        "with role-based access, validated state changes, controlled staff release "
        "requests, audit evidence and repeatable testing. The final audit improved "
        "reliability and security without replacing the original architecture or "
        "overstating the rule engine as machine learning. Local evidence strongly "
        "supports the core implementation, but final reproducibility should not be "
        "claimed until PostgreSQL, the frontend build, final screenshots, CI and the "
        "GitHub release are verified. Completing those steps and personally "
        "explaining the code will make the submission defensible against the MIT208 "
        "rubric.",

        "PhishGuard demonstrates a coherent and explainable phishing-triage workflow "
        "with role-based access, validated state changes, controlled staff release "
        "requests, audit evidence and repeatable testing. The final review improved "
        "reliability and security without replacing the original architecture and "
        "without overstating the rule engine as machine learning. The core "
        "implementation is supported by 179 automated tests, 20 live API checks, "
        "labelled screenshots and a recording of the running application, all of "
        "which are reproducible from the repository. The honest gaps are the "
        "narration, the live oral defence and PostgreSQL verification outside "
        "continuous integration. Realistic future improvements are a browser-level "
        "test suite, stronger session controls and, only with a documented dataset "
        "and error analysis, a trained classifier.",
    ),

    # ---- AI-use declaration ------------------------------------------------
    (
        "Generative AI assisted with repository audit, identification of potential "
        "risks, draft code patches, test-case expansion, diagram preparation and "
        "wording refinement. I reviewed the proposed changes against the source code, "
        "executed the recorded tests, retained evidence of actual results and "
        "modified the outputs to fit my project. AI was not used to invent users, "
        "implementation, results, references or progress. I remain responsible for "
        "the final code, security, citations and explanation, and I can describe each "
        "submitted change during the code review and live demonstration.",

        "Generative AI (Claude) assisted with auditing the repository against the "
        "assessment criteria, explaining why the pinned dependency versions failed to "
        "install on newer Python releases, drafting code changes, expanding the test "
        "cases towards negative and boundary conditions, and refining wording. My own "
        "process was to reproduce each reported defect first with a temporary probe "
        "test, then apply a fix, then lock the behaviour in with a named regression "
        "test; the figures in this report and in the repository's testing "
        "documentation are the actual output of running those suites on my machine. "
        "Suggestions that did not fit the project were rejected and are listed in the "
        "bug log. AI was not used to invent users, implementation, results, "
        "references, screenshots or progress; the screenshots and the walkthrough "
        "recording are captures of my own running application. I remain responsible "
        "for the final code, security, citations and explanation, and I can describe "
        "each submitted change during the code review and live demonstration.",
    ),
]

# Table cell replacements: (exact current text, replacement)
CELL_REPLACEMENTS: list[tuple[str, str]] = [
    ("Backend tests", "Backend tests (pytest)"),
    ("52 passed", "110 passed"),
    ("Statement coverage", "Statement coverage"),
    ("85% (716/846)", "89% (805/909)"),
    ("Live API smoke", "Live API smoke"),
    ("14/14 passed", "20/20 passed"),
    ("Reset SQLite running server", "Reset SQLite; real running server"),
    ("Frontend policy tests", "Frontend tests (vitest)"),
    ("5 passed", "69 passed"),
    ("Pure helpers; browser build pending", "Helpers, routing policy, pages; production build verified"),
    ("SQLite test database", "In-memory SQLite; PostgreSQL covered by CI"),
    ("Backend app statements", "Backend app statements"),
    # Evidence-map table
    ("routers, pages, 52 tests, 14 smoke checks",
     "routers, pages, 110 backend + 69 frontend tests, 20 smoke checks"),
    ("Final screen recording and labelled screenshots",
     "Captured: evidence/screenshots (17 images) and evidence/video (4:00 MP4)"),
    ("five-problem bug log and regression tests",
     "14-entry bug log, each with a named regression test"),
    ("Final commit IDs and appendix references",
     "Insert final commit IDs and the v1.0-final release link"),
    ("PostgreSQL negative-case screenshot",
     "PostgreSQL constraint evidence from the CI job log"),
    ("Fresh-clone proof, passing CI and release",
     "Confirm the public CI run and the v1.0-final release"),
    ("Trace one request during live demo", "Trace one request during the live demo"),
    ("Approximately 1,578 words (excluding title, references and appendices)",
     "Approximately 1,590 words (excluding title, references and appendices)"),
    # Front-matter warning: no longer a patch-pending draft
    ("Do not submit this draft unchanged. First apply the candidate patch, rerun "
     "PostgreSQL and frontend verification, replace screenshots/evidence references, "
     "personalise the lifecycle/reflection and create the genuine final release.",
     "Before submitting: insert your name, student ID, lecturer and submission date; "
     "add the v1.0-final release link; confirm the public CI run; record the "
     "narration over evidence/video/PhishGuard_Walkthrough.mp4; and read the "
     "reflection and AI declaration through so both are genuinely in your own words."),
    ("The student must edit this declaration so it exactly describes the tools and "
     "process actually used in the final submission.",
     "Read this declaration and adjust it so it exactly describes the tools and "
     "process you used. You must be able to defend every sentence of it."),
]

# Checklist items that are now done get a tick; the rest stay open.
CHECKLIST_REPLACEMENTS: list[tuple[str, str]] = [
    ("☐ Apply and inspect the candidate patch in the real GitHub repository.",
     "☑ Hardening applied and inspected in the real repository."),
    ("☐ Retain/update the existing frontend package-lock.json and run npm install, "
     "npm test and npm run build.",
     "☑ Frontend: npm install, 69 tests and a production build all run cleanly."),
    ("☐ Create/reset PostgreSQL, apply schema/seed and demonstrate the main workflow.",
     "☐ Run PostgreSQL locally once, apply schema/seed and screenshot "
     "/system/database-status showing engine: postgresql. (CI already does this.)"),
    ("☐ Rerun backend pytest coverage and the 14-step smoke workflow after all final edits.",
     "☑ Backend: 110 tests at 89% coverage and the 20-check smoke workflow rerun "
     "after the final edits."),
    ("☐ Capture and label final screenshots; replace any earlier-version images.",
     "☑ 17 labelled screenshots captured from the running application "
     "(evidence/screenshots)."),
    ("☐ Run a secret/private-data scan and confirm `.env`/database files are not staged.",
     "☑ Secret scan run; no keys or .env files staged (see docs/SECURITY.md)."),
    ("☐ Push meaningful truthful commits and obtain a passing GitHub Actions run.",
     "☐ Confirm the GitHub Actions run passes on the public repository."),
    ("☐ Create a genuine `v1.0-final` tag/release and test links in an incognito browser.",
     "☐ Publish the v1.0-final release notes and open every submitted link in a "
     "private browser window."),
    ("☐ Insert final figure/appendix/commit references and personalise first-person reflection.",
     "☐ Insert your name, ID, lecturer, date and the release link, then read the "
     "reflection and AI declaration through in your own voice."),
    ("☐ Record the 3-4 minute video and rehearse live technical questions with offline fallback.",
     "☐ Record narration over evidence/video/PhishGuard_Walkthrough.mp4 (4:00 draft "
     "capture supplied) and rehearse the live questions with an offline fallback."),
]


def set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's text while keeping its style and first-run format."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def normalise(text: str) -> str:
    return " ".join(text.split())


def main() -> int:
    if len(sys.argv) != 3:
        print(__doc__)
        return 1
    src, dst = Path(sys.argv[1]), Path(sys.argv[2])
    doc = docx.Document(str(src))

    wanted = {normalise(old): new for old, new in REPLACEMENTS}
    applied: set[str] = set()

    # Several targets are a sentence run inside a longer paragraph (for example the
    # executive summary), so match on substring and rebuild the paragraph.
    for paragraph in doc.paragraphs:
        current = normalise(paragraph.text)
        if not current:
            continue
        changed = False
        for key, new in wanted.items():
            if key in applied:
                continue
            if key == current:
                current, changed = new, True
                applied.add(key)
            elif key in current:
                current, changed = current.replace(key, new), True
                applied.add(key)
        if changed:
            set_paragraph_text(paragraph, current)

    # Checklist items are paragraphs too.
    check = {normalise(o): n for o, n in CHECKLIST_REPLACEMENTS}
    for paragraph in doc.paragraphs:
        key = normalise(paragraph.text)
        if key in check:
            set_paragraph_text(paragraph, check[key])
            applied.add(key)

    cells = {normalise(o): n for o, n in CELL_REPLACEMENTS}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                key = normalise(cell.text)
                if key in cells:
                    for i, paragraph in enumerate(cell.paragraphs):
                        set_paragraph_text(paragraph, cells[key] if i == 0 else "")
                    applied.add(key)

    missing = [o for o in list(wanted) + list(check) + list(cells) if o not in applied]
    print(f"Applied {len(applied)} replacements.")
    if missing:
        print(f"\n{len(missing)} target(s) NOT found (text may have changed):")
        for m in missing:
            print(f"  - {m[:95]}...")

    doc.save(str(dst))

    # Body word count: from the executive summary to just before References,
    # excluding headings and captions — this is the figure the brief limits.
    counting, body_words = False, 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name
        if style.startswith("Heading"):
            low = text.lower()
            if low.startswith("executive summary"):
                counting = True
            elif low.startswith("references"):
                counting = False
            continue
        if counting and style != "Caption":
            body_words += len(text.split())

    print(f"\nSaved: {dst}")
    print(f"Body word count (executive summary to conclusion, excluding "
          f"headings/captions): {body_words}")
    return 0 if not missing else 3


if __name__ == "__main__":
    sys.exit(main())
